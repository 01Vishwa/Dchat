import os
import tempfile
from docx import Document
from datetime import datetime

# Use a temp directory for exports so files are cleaned up by the OS
EXPORT_DIR = os.path.join(tempfile.gettempdir(), "dchat_exports")
os.makedirs(EXPORT_DIR, exist_ok=True)

def generate_docx(run_data: dict, qa_pairs: list) -> str:
    doc = Document()
    
    # Title & Metadata
    doc.add_heading("Vendor Security Assessment — DChat", level=1)
    
    filename = run_data.get('questionnaire_filename', 'Unknown')
    date_str = run_data.get('created_at', datetime.now().isoformat())[:10]
    doc.add_paragraph(f"Questionnaire: {filename}")
    doc.add_paragraph(f"Generated: {date_str}")
    
    # Summary
    doc.add_heading("Summary", level=2)
    total = run_data.get('total_questions', 0)
    answered = run_data.get('answered_count', 0)
    not_found = run_data.get('not_found_count', 0)
    doc.add_paragraph(
        f"Total Questions: {total} | "
        f"Answered: {answered} | "
        f"Not Found: {not_found}"
    )
    
    doc.add_paragraph()  # Spacer
    
    # Q/A Pairs
    for qa in qa_pairs:
        q_num = qa.get('question_number', '?')
        q_text = qa.get('question_text', '')
        doc.add_heading(f"Q{q_num}: {q_text}", level=2)
        
        # Prefer edited answer if available
        answer = qa.get('edited_answer') or qa.get('generated_answer', 'No answer generated.')
        doc.add_paragraph(f"Answer: {answer}")
        
        # Citations
        citations = qa.get('citations', [])
        if citations:
            doc.add_paragraph(f"Citations: {', '.join(citations)}")
            
        # Confidence Level
        conf = qa.get('confidence', 0)
        conf_level = "High" if conf > 0.7 else "Medium" if conf > 0.4 else "Low"
        conf_pct = int(conf * 100)
        doc.add_paragraph(f"Confidence: {conf_level} ({conf_pct}%)")
        
        doc.add_paragraph("─" * 40)
        
    # Save document in temp directory
    path = os.path.join(EXPORT_DIR, f"export_{run_data['id']}.docx")
    doc.save(path)
    return path
